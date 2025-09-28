import PyPDF2
import docx
import re
import os
import logging
from io import BytesIO
import pdfplumber
from typing import Optional, Dict, Any

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class ResumeParser:
    """
    Enhanced Resume Parser for PDF and DOCX files
    Uses multiple extraction methods for maximum reliability
    """

    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc']
        logger.info("ResumeParser initialized successfully")

    def extract_text(self, uploaded_file) -> Optional[str]:
        """
        Main method to extract text from uploaded file with enhanced error handling
        """
        try:
            # Handle different input types
            if hasattr(uploaded_file, 'name'):  # Streamlit file object
                file_name = uploaded_file.name
                file_content = uploaded_file.read()
                uploaded_file.seek(0)  # Reset file pointer
            else:  # File path string
                file_name = uploaded_file
                with open(uploaded_file, 'rb') as f:
                    file_content = f.read()

            # Determine file type
            file_extension = os.path.splitext(file_name)[1].lower()

            logger.info(f"Processing file: {file_name} ({len(file_content)} bytes)")

            if file_extension == '.pdf':
                return self._extract_from_pdf_enhanced(file_content, file_name)
            elif file_extension in ['.docx', '.doc']:
                return self._extract_from_docx(file_content, file_name)
            else:
                logger.error(f"Unsupported file format: {file_extension}")
                return None

        except Exception as e:
            logger.error(f"Error in extract_text: {str(e)}")
            return None

    def _extract_from_pdf_enhanced(self, file_content: bytes, file_name: str) -> Optional[str]:
        """
        Enhanced PDF extraction with multiple fallback methods
        """
        text = ""

        try:
            # Method 1: pdfplumber (most reliable for complex PDFs)
            logger.info("Trying pdfplumber extraction...")

            with pdfplumber.open(BytesIO(file_content)) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                            logger.info(f"pdfplumber: Extracted {len(page_text)} chars from page {page_num + 1}")
                    except Exception as e:
                        logger.warning(f"pdfplumber failed on page {page_num + 1}: {e}")
                        continue

            # If pdfplumber got good results, return it
            if len(text.strip()) > 100:
                logger.info(f"pdfplumber: Successfully extracted {len(text)} characters")
                return self._clean_extracted_text(text)

            # Method 2: PyPDF2 fallback
            logger.info("Trying PyPDF2 extraction as fallback...")
            text = ""

            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))

            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text += page_text + "\n"
                        logger.info(f"PyPDF2: Extracted {len(page_text)} chars from page {page_num + 1}")
                except Exception as e:
                    logger.warning(f"PyPDF2 failed on page {page_num + 1}: {e}")
                    continue

            if len(text.strip()) > 50:
                logger.info(f"PyPDF2: Successfully extracted {len(text)} characters")
                return self._clean_extracted_text(text)

            # Method 3: Try character-by-character extraction if both fail
            logger.info("Trying character extraction as last resort...")
            text = self._extract_pdf_chars(file_content)

            if text and len(text.strip()) > 20:
                logger.info(f"Character extraction: Got {len(text)} characters")
                return self._clean_extracted_text(text)

            logger.error("All PDF extraction methods failed")
            return None

        except Exception as e:
            logger.error(f"Error extracting from PDF {file_name}: {str(e)}")
            return None

    def _extract_pdf_chars(self, file_content: bytes) -> str:
        """
        Character-level extraction for difficult PDFs
        """
        try:
            with pdfplumber.open(BytesIO(file_content)) as pdf:
                text = ""
                for page in pdf.pages:
                    # Try to get characters and reconstruct text
                    chars = page.chars
                    if chars:
                        page_text = ''.join([char.get('text', '') for char in chars])
                        text += page_text + "\n"
                return text
        except:
            return ""

    def _extract_from_docx(self, file_content: bytes, file_name: str) -> Optional[str]:
        """
        Extract text from DOCX files with enhanced handling
        """
        try:
            doc = docx.Document(BytesIO(file_content))
            text = ""

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"

            if text.strip():
                logger.info(f"Successfully extracted text from DOCX: {len(text)} characters")
                return self._clean_extracted_text(text)
            else:
                logger.warning("No text found in DOCX file")
                return None

        except Exception as e:
            logger.error(f"Error extracting from DOCX {file_name}: {str(e)}")
            return None

    def _clean_extracted_text(self, text: str) -> str:
        """
        Clean and normalize extracted text with better preservation
        """
        if not text:
            return ""

        # Remove excessive whitespace but preserve structure
        text = re.sub(r'\n{3,}', '\n\n', text)  # Max 2 consecutive newlines
        text = re.sub(r'[ \t]{2,}', ' ', text)    # Multiple spaces/tabs to single space

        # Remove common OCR artifacts but preserve useful characters
        text = re.sub(r'[|•◦▪▫]{2,}', ' ', text)   # Multiple bullet points

        # Remove excessive special characters but keep formatting
        text = re.sub(r'[_-]{4,}', '---', text)    # Long underscores/dashes

        # Clean up line breaks around specific patterns
        text = re.sub(r'\s*\n\s*', '\n', text)

        return text.strip()

    def extract_contact_info(self, text: str) -> Dict[str, Any]:
        """
        Extract contact information from resume text
        """
        contact_info = {}

        if not text:
            return contact_info

        # Extract email addresses
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text, re.IGNORECASE)
        contact_info['emails'] = list(set(emails))  # Remove duplicates

        # Extract phone numbers (multiple patterns)
        phone_patterns = [
            r'\+?91[-\s]?[6-9]\d{9}',              # Indian mobile numbers
            r'\(?\+?91\)?[-\s]?\d{10}',          # Indian with country code
            r'\b(?:\+?1[-.]?)?(?:\(?[0-9]{3}\)?[-.]?)?[0-9]{3}[-.]?[0-9]{4}\b',  # US format
            r'\(\d{3}\)\s*\d{3}-\d{4}',         # (123) 456-7890
            r'\d{3}[-.]\d{3}[-.]\d{4}',            # 123-456-7890
            r'\+?[1-9]\d{7,14}'                     # International format
        ]

        phones = []
        for pattern in phone_patterns:
            matches = re.findall(pattern, text)
            phones.extend(matches)

        # Clean and validate phone numbers
        cleaned_phones = []
        for phone in phones:
            # Remove formatting and keep only digits and +
            clean_phone = re.sub(r'[^\d+]', '', phone)
            if 10 <= len(clean_phone.replace('+', '')) <= 15:  # Valid phone number length
                cleaned_phones.append(phone)

        contact_info['phones'] = list(set(cleaned_phones))

        # Extract LinkedIn profiles
        linkedin_pattern = r'linkedin\.com/in/[A-Za-z0-9-]+|linkedin\.com/pub/[A-Za-z0-9-/]+'
        linkedin_profiles = re.findall(linkedin_pattern, text, re.IGNORECASE)
        contact_info['linkedin'] = list(set(linkedin_profiles))

        return contact_info

    def validate_resume(self, text: str) -> Dict[str, Any]:
        """
        Validate if the extracted text looks like a resume
        """
        validation = {
            'is_valid': False,
            'confidence': 0.0,
            'issues': [],
            'strengths': []
        }

        if not text or len(text.strip()) < 50:
            validation['issues'].append("Text too short to be a resume")
            return validation

        score = 0
        max_score = 10

        # Check for resume-like keywords
        resume_keywords = [
            'experience', 'education', 'skills', 'work', 'university', 'college',
            'degree', 'bachelor', 'master', 'phd', 'certified', 'project',
            'company', 'position', 'role', 'responsibility', 'achievement',
            'internship', 'training', 'course', 'programming', 'software'
        ]

        text_lower = text.lower()
        found_keywords = [kw for kw in resume_keywords if kw in text_lower]

        if len(found_keywords) >= 8:
            score += 4
            validation['strengths'].append(f"Contains {len(found_keywords)} resume keywords")
        elif len(found_keywords) >= 5:
            score += 3
        elif len(found_keywords) >= 3:
            score += 2
        else:
            validation['issues'].append("Few resume-related keywords found")

        # Check for contact information
        contact_info = self.extract_contact_info(text)
        if contact_info.get('emails'):
            score += 2
            validation['strengths'].append("Contains email address")
        else:
            validation['issues'].append("No email address found")

        if contact_info.get('phones'):
            score += 1
            validation['strengths'].append("Contains phone number")

        # Check text length (resumes should be reasonable length)
        word_count = len(text.split())
        if 100 <= word_count <= 3000:
            score += 2
            validation['strengths'].append(f"Appropriate length ({word_count} words)")
        elif word_count < 100:
            validation['issues'].append("Text seems too short for a resume")
        else:
            validation['issues'].append("Text seems too long for a resume")

        # Check for sections
        section_keywords = ['education', 'experience', 'skills', 'project']
        sections_found = sum(1 for kw in section_keywords if kw in text_lower)

        if sections_found >= 3:
            score += 1
            validation['strengths'].append(f"Contains {sections_found} resume sections")
        else:
            validation['issues'].append("Few identifiable resume sections")

        validation['confidence'] = score / max_score
        validation['is_valid'] = validation['confidence'] >= 0.5

        return validation

# Test function
def test_parser():
    """Test the resume parser with sample text"""
    parser = ResumeParser()

    # Test text (simulate extracted content)
    sample_text = """
    ABHAY RUDATALA
    CONTACT
    +91 9687793375
    abhayrudatala56789@gmail.com
    Ahmedabad, Gujarat, 380009

    TECHNICAL SKILLS
    Languages: C/C++, Python, JavaScript
    Frameworks: Django, Express.js

    EDUCATION
    Master of Computer Applications
    GLS UNIVERSITY (2024-PRESENT)
    Current CGPA: 8.36
    """

    # Test contact extraction
    contact_info = parser.extract_contact_info(sample_text)
    print("Contact Info:", contact_info)

    # Test validation
    validation = parser.validate_resume(sample_text)
    print("Validation:", validation)

if __name__ == "__main__":
    test_parser()